try:
    from docx import Document
    import sys

    def read_docx(file_path):
        doc = Document(file_path)
        full_text = []
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return '\n'.join(full_text)

    if __name__ == "__main__":
        if len(sys.argv) > 1:
            print(read_docx(sys.argv[1]))
        else:
            print("Please provide a file path")
except Exception as e:
    print(f"Error: {e}")
